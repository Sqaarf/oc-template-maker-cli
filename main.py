import matplotlib.pyplot as plt
import pandas as pd
from math import pi

import docx

def parse_conf():
    content = []
    with open("./template.conf", "r", encoding="utf-8") as f:
        for line in f.readlines():
            content.append(line.strip())
    return content

def generate_radar_chart(data_frame):
    categories=list(data_frame)[1:]
    N = len(categories)
    values=data_frame.loc[0].drop('group').values.flatten().tolist()
    values += values[:1]

    angles = [n / float(N) * 2 * pi for n in range(N)]
    angles += angles[:1]
    
    ax = plt.subplot(111, polar=True)
    
    plt.xticks(angles[:-1], categories, color='grey', size=8)
    
    ax.set_rlabel_position(0)
    plt.yticks([x for x in range(10)], [str(x) for x in range(10)], color="grey", size=7)
    plt.ylim(0,10)
    
    ax.plot(angles, values, 'b', linewidth=1, linestyle='solid')
    ax.fill(angles, values, 'b', alpha=0.1)
    plt.savefig('skills.png')


content_conf_file = parse_conf()
main_title, entries = (content_conf_file[0], content_conf_file[1:])
skill = False

document = docx.Document()

document.add_heading(main_title, 0)

for entry in entries:
    entry_title, entry_content = entry.split(":")
    if entry_title.lower() == "!skills":
        skill = True
        skills_data = tuple(map(int, entry_content.split(",")))
        df = pd.DataFrame({
            'group': ['A'],
            'STR': skills_data[0],
            'SPD': skills_data[1],
            'STA': skills_data[2],
            'RES': skills_data[3],
            'MAG': skills_data[4],
            'CHA': skills_data[5],
            'INT': skills_data[6]
            })
        generate_radar_chart(df)
        document.add_picture('./skills.png')
    else:
        document.add_heading(entry_title, level=1)
        p = document.add_paragraph(entry_content)


document.save('demo.docx')